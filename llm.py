from typing import List
from schema import Relationship, RelationshipList, RelationshipLite, RelationshipLiteList
from utils import convert_to_lite, get_dataframe, df2json, get_unique_entities
from json_data import sample_results
import pandas as pd
import streamlit as st
import json
from langchain_openai import ChatOpenAI
from langchain_experimental.graph_transformers import LLMGraphTransformer
from langchain_core.documents import Document
import json
from typing import List, Dict
import PyPDF2
import docx
import io
from prompts import ENTITY_EXTRACTION_PROMPT
from config import LLM_CONFIG


def read_pdf(file) -> str:
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def read_docx(file) -> str:
    doc = docx.Document(io.BytesIO(file.read()))
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def read_txt(file) -> str:
    return file.getvalue().decode("utf-8")

def extract_content(file) -> str:
    file_extension = file.name.split('.')[-1].lower()
    
    if file_extension == 'pdf':
        return read_pdf(file)
    elif file_extension == 'docx':
        return read_docx(file)
    elif file_extension == 'txt':
        return read_txt(file)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

# Function that pretends to process the documents and extract the relationships but returns a sample result from json.py
def process_documents_sample(files: List) -> List[RelationshipLite]:
    
    # Wrap the list in a dictionary with 'relationships' key
    if isinstance(sample_results, list):
        validated_data = RelationshipList(relationships=sample_results)
    else:
        # If response is already in the expected format
        validated_data = RelationshipList(**sample_results)

    # Convert the relationships to a list of RelationshipLite
    lite_results = convert_to_lite(validated_data.relationships)
    return lite_results

# Process the documents and extract the relationships
def process_documents(files: List) -> List[RelationshipLite]:
    llm = ChatOpenAI(
        model=LLM_CONFIG["model"],
        temperature=LLM_CONFIG["temperature"]
    )

    all_results = []
    for file in files:
        content = extract_content(file)
        prompt = ENTITY_EXTRACTION_PROMPT.format(content=content)
        response = llm.invoke(prompt)

        
        try:
            # Handle different response types
            if hasattr(response, 'content'):
                response_text = response.content
            elif isinstance(response, str):
                response_text = response
            else:
                response_text = str(response)
            
            # Try to clean the response text
            response_text = response_text.strip()
            if not response_text:
                st.error(f"Empty response from LLM for file: {file.name}")
                continue

            # If the response_text starts with the word ```json in the beginning, remove it
            if response_text.startswith('```json'):
                response_text = response_text[7:]

            # If the response_text ends with the word ``` in the end, remove it
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            
            # Parse the raw JSON response
            raw_data = json.loads(response_text)
            
            # Wrap the list in a dictionary with 'relationships' key
            if isinstance(raw_data, list):
                validated_data = RelationshipList(relationships=raw_data)
            else:
                # If response is already in the expected format
                validated_data = RelationshipList(**raw_data)

            # Convert the relationships to a list of RelationshipLite
            lite_results = convert_to_lite(validated_data.relationships)
            all_results.extend(lite_results)
        except json.JSONDecodeError as e:
            st.error(f"Failed to parse LLM response for file: {file.name}")
            st.error(f"JSON Error: {str(e)}")
            st.error(f"Response text: {response_text}")
            continue
        except ValueError as e:
            st.error(f"Invalid relationship structure in response for file {file.name}: {str(e)}")
            continue
        except Exception as e:
            st.error(f"Unexpected error processing file {file.name}: {str(e)}")
            continue
            
    return all_results

def extract_graph(files: List, edited_df: pd.DataFrame) -> list[Dict]:
    # Create an empty list to store the graphs
    graphs = []

    # Initialize the LLM
    llm = ChatOpenAI(
        model=LLM_CONFIG["model"],
        temperature=LLM_CONFIG["temperature"]
    )

    allowed_relationships = df2json(edited_df)
    allowed_nodes = get_unique_entities(edited_df)

    print(f"Allowed Relationships: {allowed_relationships}")
    print(f"Allowed Nodes: {allowed_nodes}")    

    # Create a graph transformer
    graph_transformer = LLMGraphTransformer(llm=llm, 
                                            allowed_relationships=allowed_relationships, 
                                            allowed_nodes=allowed_nodes, 
                                            node_properties=True, 
                                            relationship_properties=True)

    for file in files:
        content = extract_content(file)
        documents = [Document(page_content=content)]
        data = graph_transformer.convert_to_graph_documents(documents)
        
        graphs.append(data)
        print("-"*100)
        print(f"Nodes:{data[0].nodes}")
        print("-"*100)
        print(f"Relationships:{data[0].relationships}")
        print("-"*100)

    return graphs


        
